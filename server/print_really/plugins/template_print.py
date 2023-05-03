from flask import Flask, jsonify, request
from python_docx_replace import docx_replace
from docx import Document
from config.configs import AUTH_key
import requests
import random
import string
import re

from config.configs import localhost_path

def get_file_path(dayi_cookie,uuid):
    cookies = {'dayi-cookie-for-uploads':dayi_cookie}
    r = requests.get(localhost_path+"/api/get_user_files",cookies=cookies)
    if r.json()['code'] != 201 :
        return 'Error' # No login
    for l in r.json()['data']:
        if l['uuid'] == uuid:
            return l['file_path']

# 插件注册
def register_plugin(app):
    # 获取模板中的可替换项并传递向前端
    # 可替换的点位用这种方式表示：
    # ${text}
    # 这样替换时只需要在POST的json里面添加上该项就行了，详细格式见下面
    # 使用GET传输，参数file为文件的uuid
    # 返回的数据样例：
    # {code:"201",info:"已找到替换点",data:["aaa","bbb","ccc","这里是可以替换的点位"]，text:'这里返回全文的文本内容,方便前端简单的展示一下'}
    # 模板文件建议规定文件后缀方便区分，比如{name}.template.docx
  @app.route(AUTH_key+"/api-v2/get_template", methods=['GET'])
  def get_template():
    file_uuid = request.args.get('file')
    file_path = get_file_path(request.cookies.get('dayi-cookie-for-uploads'),file_uuid)
    if file_path == 'Error':
        return jsonify({'code': 412,'info':'打开文件错误'})
    if file_path.endswith('.template.docx') == False:
        return jsonify({'code': 413,'info':'文件不是模板文件'})
    doc_str = ''
    doc = Document(file_path)
    for p in doc.paragraphs:
        doc_str = doc_str + '\n' + p.text
    rep_point = list(set(re.findall(r"[$][{](.*?)[}]",doc_str)))
    return jsonify({'code':201,'info':'已找到替换点','data':rep_point,'text':doc_str})
    pass
    # 进行替换
    # 表单样例
    # {code:"201",info:"前端传递的数据",data:{file:"rhsdfgh-fgshfgh-dfgdfsg-dghfdgh",replace:{name:"hihihi",date:"1145/1/4"}}}
    # 上面要的还是文件的uuid
    # 模板文件建议规定文件后缀方便区分，比如{name}.template.docx
    # 要docx格式的文档，注意
  @app.route(AUTH_key+"/api-v2/template_replace", methods=['POST'])
  def template_replace():
    json_data = request.json
    if int(json_data['code']) != 201:
      return jsonify({'code': 411,'info':'非法参数'})
    file_uuid = json_data['data']['file']
    file_path = get_file_path(request.cookies.get('dayi-cookie-for-uploads'),file_uuid)
    if file_path == 'Error':
        return jsonify({'code': 412,'info':'打开文件错误'})
    if file_path.endswith('.template.docx') == False:
        return jsonify({'code': 413,'info':'文件不是模板文件'})
    doc = Document(file_path)
    if type(json_data['data']['replace']) != dict:
        return jsonify({'code': 414,'info':'格式错误'})
    docx_replace(doc,**json_data['data']['replace'])
    new_path = file_path+''.join(random.choice(string.ascii_lowercase) for i in range(16))+'.docx'
    doc.save(new_path)
    # !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
    # !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
    # !!!!!!!! 这里需要一个向后端数据库注册文件的代码 !!!!!!!!!!!
    # !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
    # !!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!
    return jsonify({'code': 201,'info':'成功'})

#加载插件
def load_plugin(app):
  register_plugin(app)